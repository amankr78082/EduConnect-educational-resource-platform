from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "docs/RGPV_Cloud_Computing_Unit_I_Notes.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx] / 1440)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_font(run, name="Calibri", size=11, bold=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.10
    if bold_prefix and text.startswith(bold_prefix):
        r = p.add_run(bold_prefix)
        set_font(r, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_font(r2)
    else:
        r = p.add_run(text)
        set_font(r)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_font(r)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.10
    r = p.add_run(text)
    set_font(r)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    color = "2E74B5" if level in (1, 2) else "1F4D78"
    size = {1: 16, 2: 13, 3: 12}.get(level, 12)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    for run in p.runs:
        set_font(run, size=size, bold=True, color=color)
    return p


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        set_font(r, size=10, bold=True)
        set_cell_shading(hdr[i], "E8EEF5")
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(value)
            set_font(r, size=10)
    set_table_width(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_answer_block(doc, title, intro, points, example, conclusion):
    add_heading(doc, title, 2)
    add_para(doc, "Exam answer opening: " + intro, bold_prefix="Exam answer opening:")
    add_para(doc, "Main points:")
    for point in points:
        add_bullet(doc, point)
    add_para(doc, "Real-life example: " + example, bold_prefix="Real-life example:")
    add_para(doc, "Conclusion: " + conclusion, bold_prefix="Conclusion:")


def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ["top_margin", "bottom_margin", "left_margin", "right_margin"]:
        setattr(section, side, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color in [
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        st._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor.from_string(color)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("RGPV Cloud Computing - Unit I Notes")
    set_font(r, size=9, color="666666")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(3)
    r = title.add_run("Cloud Computing Fundamentals")
    set_font(r, size=24, bold=True, color="0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("RGPV Exam Notes - Unit I | 10-Mark Answer Format with Examples")
    set_font(r, size=12, color="555555")
    add_para(doc, "How to use these notes: Learn the headings, write answers in definition -> explanation -> diagram/table -> example -> conclusion order, and add the real-life example to make the answer look complete.")

    add_heading(doc, "Unit I Syllabus Coverage", 1)
    syllabus = [
        "Cloud Computing definition",
        "Types of cloud",
        "Cloud services",
        "Benefits and challenges of cloud computing",
        "Evolution of Cloud Computing",
        "NIST architecture of cloud computing",
        "Applications of cloud computing",
        "Business models around cloud",
        "Major players/tools: Eucalyptus, Nimbus, OpenNebula, CloudSim, VMware",
    ]
    for item in syllabus:
        add_bullet(doc, item)

    add_heading(doc, "1. Cloud Computing Definition", 1)
    add_answer_block(
        doc,
        "10-Mark Explanation: What is Cloud Computing?",
        "Cloud computing is a model of delivering computing resources such as servers, storage, databases, networking, software, analytics, and applications over the Internet on demand, with pay-as-you-use pricing and elastic scalability.",
        [
            "In traditional computing, an organization buys and maintains its own hardware. In cloud computing, the organization rents resources from a cloud provider.",
            "The resources are available through a network, usually the Internet, and users can access them from laptops, mobiles, tablets, or enterprise systems.",
            "The cloud provider manages the physical infrastructure, power, cooling, hardware maintenance, virtualization, and many security controls.",
            "Users can increase or decrease resources based on demand. This property is called elasticity.",
            "Cloud computing supports measured service, so users pay according to usage such as storage consumed, virtual machines used, or data transferred.",
        ],
        "When a student uploads photos to Google Drive, they are not storing data only on their phone. The files are stored in Google's cloud data centers and can be accessed from any device after login.",
        "Cloud computing converts computing from a product owned by the user into a service consumed on demand."
    )
    add_table(
        doc,
        ["Characteristic", "Meaning", "Example"],
        [
            ["On-demand self-service", "User can provision resources without manual provider interaction.", "Create a virtual server from AWS or Azure portal."],
            ["Broad network access", "Services are available over a network through standard devices.", "Access Gmail from mobile and laptop."],
            ["Resource pooling", "Provider shares resources among multiple users using multi-tenancy.", "Many customers use the same physical data center."],
            ["Rapid elasticity", "Resources can scale up or down quickly.", "E-commerce site adds servers during a sale."],
            ["Measured service", "Usage is monitored and billed.", "Pay for GB storage used per month."],
        ],
        [2100, 4200, 3060],
    )

    add_heading(doc, "2. Types of Cloud", 1)
    add_para(doc, "Cloud deployment models describe who owns the cloud, who can access it, and how it is managed.")
    add_table(
        doc,
        ["Type", "Description", "Real-Life Example", "Best Used For"],
        [
            ["Public cloud", "Owned by third-party provider and shared among many customers.", "AWS, Microsoft Azure, Google Cloud.", "Startups, web apps, scalable services."],
            ["Private cloud", "Dedicated cloud for one organization, managed internally or by a vendor.", "A bank's internal cloud for sensitive transactions.", "High control, compliance, sensitive data."],
            ["Hybrid cloud", "Combination of private and public cloud with data/app portability.", "Hospital keeps patient records private and uses public cloud for analytics.", "Balancing security and scalability."],
            ["Community cloud", "Shared by organizations with common requirements.", "Cloud shared by government departments or universities.", "Common compliance, cost sharing."],
        ],
        [1700, 2900, 2500, 2260],
    )
    add_answer_block(
        doc,
        "10-Mark Explanation: Public, Private, Hybrid, and Community Cloud",
        "The main types of cloud are public, private, hybrid, and community cloud. They differ based on ownership, access, control, cost, and security.",
        [
            "Public cloud provides low entry cost and high scalability but offers less direct control over infrastructure.",
            "Private cloud gives better control, customization, and privacy, but it is expensive to build and maintain.",
            "Hybrid cloud gives flexibility by keeping sensitive workloads in private cloud and using public cloud for extra capacity.",
            "Community cloud is useful when multiple organizations have similar security, policy, or mission requirements.",
            "The choice depends on budget, compliance, performance, and data sensitivity.",
        ],
        "An online coaching platform can run video classes on a public cloud, while a bank may use private cloud for core banking and public cloud only for marketing analytics.",
        "Cloud deployment models help organizations choose the right balance between cost, scalability, control, and security."
    )

    add_heading(doc, "3. Cloud Services", 1)
    add_para(doc, "Cloud service models describe how much responsibility is managed by the provider and how much is managed by the customer.")
    add_table(
        doc,
        ["Service Model", "Provider Gives", "User Manages", "Example"],
        [
            ["IaaS", "Virtual machines, storage, networking, basic infrastructure.", "OS, middleware, runtime, application, data.", "Amazon EC2, Azure Virtual Machines."],
            ["PaaS", "Platform, runtime, database tools, deployment environment.", "Application code and data.", "Google App Engine, Heroku."],
            ["SaaS", "Complete ready-to-use software application.", "Only user settings and data usage.", "Gmail, Salesforce, Microsoft 365."],
        ],
        [1600, 3100, 2500, 2160],
    )
    add_answer_block(
        doc,
        "10-Mark Explanation: IaaS, PaaS, and SaaS",
        "Cloud services are commonly divided into Infrastructure as a Service, Platform as a Service, and Software as a Service. These models reduce the burden of ownership at different levels.",
        [
            "IaaS provides basic computing infrastructure such as virtual servers, storage, and networks. It gives high control but requires technical management.",
            "PaaS provides a development and deployment platform. Developers focus on code while the provider manages servers, OS, and runtime.",
            "SaaS provides complete applications through a browser or app. Users do not manage infrastructure or installation.",
            "IaaS is suitable for system administrators and enterprises needing flexibility; PaaS is suitable for developers; SaaS is suitable for end users.",
            "As we move from IaaS to SaaS, user control decreases but convenience increases.",
        ],
        "For an online food delivery startup, IaaS can host custom servers, PaaS can deploy the app quickly, and SaaS such as Gmail or Zoho can handle company email and CRM.",
        "Service models make cloud computing useful for different users, from infrastructure engineers to normal office workers."
    )

    add_heading(doc, "4. Benefits and Challenges of Cloud Computing", 1)
    add_table(
        doc,
        ["Benefits", "Explanation", "Example"],
        [
            ["Cost saving", "No need to buy servers in advance; pay according to use.", "Startup launches app without owning data center."],
            ["Scalability", "Resources increase or decrease with demand.", "Shopping website handles Diwali sale traffic."],
            ["Accessibility", "Access applications and data from anywhere.", "Students access notes from phone and laptop."],
            ["Disaster recovery", "Data can be backed up across regions.", "Company restores data after local system failure."],
            ["Faster deployment", "Servers and services can be created in minutes.", "Developer tests a new app quickly."],
        ],
        [1800, 4300, 3260],
    )
    add_table(
        doc,
        ["Challenges", "Explanation", "Example"],
        [
            ["Security and privacy", "Data is stored outside the direct physical control of the user.", "Customer data must be protected from unauthorized access."],
            ["Downtime", "Internet or provider outage can affect services.", "A cloud email outage stops office communication."],
            ["Vendor lock-in", "Moving from one provider to another can be difficult.", "App uses provider-specific database services."],
            ["Compliance", "Laws may require data to stay in a country or follow audit rules.", "Banking and healthcare data need strict compliance."],
            ["Cost control", "Poor monitoring can increase bills.", "Unused virtual machines continue running."],
        ],
        [1800, 4300, 3260],
    )
    add_answer_block(
        doc,
        "10-Mark Explanation: Benefits and Challenges",
        "Cloud computing provides flexibility, cost efficiency, and speed, but it also introduces security, dependency, and management challenges.",
        [
            "The biggest benefit is reduced capital expenditure because companies rent resources instead of buying servers.",
            "Cloud improves business continuity through backups, replication, and disaster recovery.",
            "It supports remote work because applications and data are available through the Internet.",
            "Main challenges include data security, privacy, legal compliance, vendor lock-in, downtime, and cost monitoring.",
            "Good cloud governance, encryption, identity management, backups, and monitoring reduce these challenges.",
        ],
        "Netflix uses cloud infrastructure to stream content globally and scale during peak hours. At the same time, it must continuously manage security, performance, and cost.",
        "Cloud computing is powerful, but it must be planned and governed properly to get maximum benefit."
    )

    add_heading(doc, "5. Evolution of Cloud Computing", 1)
    add_para(doc, "Cloud computing did not appear suddenly. It evolved from earlier computing models.")
    add_table(
        doc,
        ["Stage", "Main Idea", "Contribution to Cloud"],
        [
            ["Mainframe computing", "Many users shared one large central computer.", "Introduced shared computing resources."],
            ["Client-server computing", "Clients requested services from dedicated servers.", "Separated user interface and backend services."],
            ["Distributed computing", "Workload divided across multiple machines.", "Improved performance and reliability."],
            ["Grid computing", "Geographically distributed resources worked together.", "Enabled large-scale resource sharing."],
            ["Virtualization", "One physical server runs multiple virtual machines.", "Foundation of elastic resource pooling."],
            ["Utility computing", "Computing charged like electricity or water.", "Introduced pay-per-use model."],
            ["Cloud computing", "Internet-based, on-demand, elastic computing services.", "Combines virtualization, automation, and service delivery."],
        ],
        [2100, 3600, 3660],
    )
    add_answer_block(
        doc,
        "10-Mark Explanation: Evolution of Cloud Computing",
        "Cloud computing evolved through mainframe sharing, client-server architecture, distributed computing, grid computing, virtualization, and utility computing.",
        [
            "Mainframes introduced the idea that many users can share one powerful computer.",
            "Client-server systems separated client devices from backend servers.",
            "Distributed and grid computing used many machines together to solve large problems.",
            "Virtualization allowed multiple virtual machines to run on one physical server, improving utilization.",
            "Utility computing introduced the business idea of charging according to usage, which became central to cloud computing.",
        ],
        "Earlier, a company had to buy a server for a website. Today, the same company can rent a virtual server, database, and storage from a cloud provider within minutes.",
        "Modern cloud computing is the result of technical progress in networking, virtualization, automation, and service-based business models."
    )

    add_heading(doc, "6. NIST Architecture of Cloud Computing", 1)
    add_para(doc, "NIST describes cloud computing using essential characteristics, service models, deployment models, and major actors.")
    add_table(
        doc,
        ["NIST Component", "Content"],
        [
            ["Essential characteristics", "On-demand self-service, broad network access, resource pooling, rapid elasticity, measured service."],
            ["Service models", "IaaS, PaaS, SaaS."],
            ["Deployment models", "Public, private, hybrid, community cloud."],
            ["Major actors", "Cloud consumer, provider, broker, auditor, carrier."],
        ],
        [2600, 6760],
    )
    add_table(
        doc,
        ["Actor", "Role"],
        [
            ["Cloud consumer", "Person or organization that uses cloud services."],
            ["Cloud provider", "Entity that offers cloud infrastructure, platform, or software services."],
            ["Cloud broker", "Manages use, performance, delivery, and relationships among providers and consumers."],
            ["Cloud auditor", "Performs independent assessment of security, performance, and compliance."],
            ["Cloud carrier", "Provides network connectivity between provider and consumer."],
        ],
        [2200, 7160],
    )
    add_answer_block(
        doc,
        "10-Mark Explanation: NIST Cloud Architecture",
        "The NIST architecture gives a standard view of cloud computing. It explains the essential characteristics, service models, deployment models, and actors involved in cloud service delivery.",
        [
            "Cloud consumer requests and uses cloud services such as storage, virtual machines, or applications.",
            "Cloud provider owns or manages the infrastructure and delivers cloud services.",
            "Cloud broker helps select, combine, or manage services from different providers.",
            "Cloud auditor checks security, privacy, performance, and compliance.",
            "Cloud carrier provides network connectivity, similar to an Internet service provider.",
        ],
        "When a college uses Microsoft 365, the college is the consumer, Microsoft is the provider, the Internet company is the carrier, and an external security agency may act as auditor.",
        "NIST architecture is important because it provides a common language for understanding cloud systems."
    )

    add_heading(doc, "7. Applications of Cloud Computing", 1)
    add_table(
        doc,
        ["Application Area", "Use of Cloud", "Real-Life Example"],
        [
            ["Education", "Online classes, digital notes, LMS, exam portals.", "Google Classroom, Moodle cloud hosting."],
            ["Healthcare", "Patient records, telemedicine, medical image storage.", "Hospital stores MRI scans in cloud storage."],
            ["Business", "ERP, CRM, email, collaboration, analytics.", "Salesforce CRM and Microsoft 365."],
            ["Entertainment", "Streaming, gaming, content delivery.", "Netflix and cloud gaming platforms."],
            ["Banking", "Mobile banking, fraud detection, data backup.", "Banks use cloud analytics for fraud patterns."],
            ["E-commerce", "Scalable web hosting, payment processing, recommendation engines.", "Amazon-like sale traffic scaling."],
            ["IoT and smart cities", "Collecting and analyzing sensor data.", "Traffic sensors send data to cloud analytics platform."],
        ],
        [1900, 4200, 3260],
    )
    add_answer_block(
        doc,
        "10-Mark Explanation: Applications of Cloud Computing",
        "Cloud computing is used in almost every field because it provides storage, processing power, software access, collaboration, and analytics through the Internet.",
        [
            "In education, cloud supports online learning, video lectures, assignments, and digital libraries.",
            "In healthcare, it supports telemedicine, patient records, and medical image storage.",
            "In business, it supports CRM, ERP, email, collaboration, and decision-making dashboards.",
            "In entertainment, it supports video streaming, music streaming, and online gaming.",
            "In IoT, cloud collects large amounts of sensor data and performs analytics.",
        ],
        "During online exams or classes, thousands of students can access the platform together because the backend can run on scalable cloud infrastructure.",
        "Cloud applications show that cloud computing is not only a technical concept but also a practical foundation for modern digital services."
    )

    add_heading(doc, "8. Business Models Around Cloud", 1)
    add_para(doc, "Cloud business models describe how providers earn revenue and how customers consume services.")
    add_table(
        doc,
        ["Business Model", "Meaning", "Example"],
        [
            ["Pay-as-you-go", "Customer pays only for actual usage.", "Pay per hour for virtual machine."],
            ["Subscription", "Fixed periodic payment for software or platform.", "Monthly Microsoft 365 subscription."],
            ["Freemium", "Basic service free; advanced features paid.", "Free storage with paid upgrade."],
            ["Reserved capacity", "Lower price for long-term commitment.", "Reserve cloud servers for one or three years."],
            ["Marketplace model", "Provider hosts third-party apps/services and earns commission.", "Cloud marketplace selling security tools."],
            ["Managed services", "Provider or partner manages cloud operations for customer.", "Cloud migration and monitoring service."],
        ],
        [2100, 4300, 2960],
    )
    add_answer_block(
        doc,
        "10-Mark Explanation: Cloud Business Models",
        "Cloud business models are based on delivering computing as a service and charging customers according to access, usage, subscription, or value-added management.",
        [
            "Pay-as-you-go is the most common model and reduces waste because customers pay only for consumed resources.",
            "Subscription model is common in SaaS, where customers pay monthly or yearly for software access.",
            "Freemium attracts users by giving basic features free and charging for extra storage, users, or advanced functions.",
            "Reserved capacity gives discounts when customers commit to long-term usage.",
            "Managed services help companies that want cloud benefits but do not have enough in-house expertise.",
        ],
        "A small business may use free cloud email initially, then buy a paid subscription as employees increase, and later use managed cloud backup for business continuity.",
        "Cloud business models make advanced technology affordable and flexible for organizations of all sizes."
    )

    add_heading(doc, "9. Major Players and Tools in Cloud Computing", 1)
    add_para(doc, "The syllabus names several cloud platforms, simulators, and virtualization tools. For exams, write their purpose and one use case.")
    add_table(
        doc,
        ["Name", "Type", "Purpose / Importance", "Example Use"],
        [
            ["Eucalyptus", "Open-source cloud platform", "Used to build private and hybrid clouds compatible with AWS-style APIs.", "University lab creates a private cloud for experiments."],
            ["Nimbus", "Cloud toolkit", "Designed for scientific and academic cloud computing using virtual machines.", "Research group runs scientific workloads."],
            ["OpenNebula", "Cloud management platform", "Manages virtualized data centers and builds private, public, and hybrid clouds.", "Company manages internal virtual infrastructure."],
            ["CloudSim", "Simulation toolkit", "Allows researchers to model and test cloud environments without real infrastructure.", "Student simulates VM scheduling algorithm."],
            ["VMware", "Virtualization and cloud software", "Provides enterprise virtualization, private cloud, and data center management tools.", "Company runs many virtual servers on fewer physical machines."],
        ],
        [1400, 1700, 4100, 2160],
    )
    add_answer_block(
        doc,
        "10-Mark Explanation: Major Players and Tools",
        "Cloud computing includes commercial providers, open-source platforms, simulation tools, and virtualization vendors. The syllabus highlights Eucalyptus, Nimbus, OpenNebula, CloudSim, and VMware.",
        [
            "Eucalyptus helps build AWS-compatible private or hybrid clouds.",
            "Nimbus focuses on scientific and academic cloud use cases.",
            "OpenNebula manages virtualized data centers and supports private, public, and hybrid cloud deployments.",
            "CloudSim is not a real cloud provider; it is a simulation toolkit for studying cloud algorithms, resource allocation, and scheduling.",
            "VMware is a major virtualization company whose technologies help create and manage virtual infrastructure, which is a foundation for private clouds.",
        ],
        "If an engineering student wants to test how virtual machines are scheduled in a data center, using real servers is costly. CloudSim allows simulation on a normal computer.",
        "These tools show the practical ecosystem of cloud computing, from real infrastructure management to research simulation."
    )

    add_heading(doc, "Quick Revision: One-Page Memory Sheet", 1)
    add_table(
        doc,
        ["Question", "Must Write Points"],
        [
            ["Definition", "Internet-based on-demand delivery of computing resources; elasticity; measured service; pay-per-use."],
            ["Types of cloud", "Public, private, hybrid, community; compare ownership, cost, security, control."],
            ["Services", "IaaS gives infrastructure; PaaS gives platform; SaaS gives complete software."],
            ["Benefits", "Cost saving, scalability, accessibility, disaster recovery, faster deployment."],
            ["Challenges", "Security, privacy, downtime, vendor lock-in, compliance, cost control."],
            ["Evolution", "Mainframe -> client-server -> distributed -> grid -> virtualization -> utility -> cloud."],
            ["NIST", "5 characteristics, 3 service models, 4 deployment models, actors: consumer, provider, broker, auditor, carrier."],
            ["Applications", "Education, healthcare, business, entertainment, banking, e-commerce, IoT."],
            ["Business models", "Pay-as-you-go, subscription, freemium, reserved capacity, marketplace, managed services."],
            ["Tools", "Eucalyptus, Nimbus, OpenNebula, CloudSim, VMware."],
        ],
        [2400, 6960],
    )

    add_heading(doc, "How to Frame Any 10-Mark Answer", 1)
    for step in [
        "Start with a two-line definition.",
        "Explain the concept in 4-6 logical points.",
        "Add a small table or architecture diagram if possible.",
        "Write one real-life example related to education, banking, e-commerce, or healthcare.",
        "End with a conclusion showing importance or practical use.",
    ]:
        add_number(doc, step)

    doc.save(OUT)


if __name__ == "__main__":
    build()
